import io
import json
import os
from datetime import datetime, timezone
from typing import Any
from urllib.parse import quote_plus

import pandas as pd
import requests
import streamlit as st
from dateutil import parser as date_parser
from docx import Document
from openai import OpenAI
from supabase import Client, create_client

st.set_page_config(page_title="GrantAI Europe", page_icon="🇪🇺", layout="wide")

EC_SEARCH_API = "https://api.tech.ec.europa.eu/search-api/prod/rest/search"
EC_API_KEY = "SEDIA"
PORTAL_BASE = "https://ec.europa.eu/info/funding-tenders/opportunities/portal/screen/opportunities/topic-search"

DEFAULT_ORG = {
    "legal_name": "II Ciobotaru Viorel Razvan Ionut",
    "country": "Romania",
    "organisation_type": "SME / întreprindere individuală",
    "pic": "",
    "caen": "0111",
    "turnover_eur": 0,
    "staff": 0,
    "capabilities": [
        "agriculture",
        "smart greenhouse",
        "renewable energy",
        "battery storage",
        "AI automation",
    ],
    "past_projects": [],
}

DEFAULT_PROJECT = {
    "name": "GreenRise",
    "summary": "Seră inteligentă cu energie regenerabilă, baterii și automatizare AI.",
    "keywords": ["agriculture", "greenhouse", "energy", "battery", "AI", "rural", "agrifood"],
    "target_budget_eur": 0,
    "preferred_role": "beneficiary",
}

PROGRAMMES = [
    "Toate",
    "Horizon Europe",
    "Digital Europe",
    "LIFE",
    "Erasmus+",
    "Connecting Europe Facility",
    "Single Market Programme",
    "EU4Health",
]

def get_secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default))
    except Exception:
        return os.getenv(name, default)

@st.cache_resource
def get_supabase() -> Client:
    url = get_secret("SUPABASE_URL")
    key = get_secret("SUPABASE_ANON_KEY")
    if not url or not key:
        raise RuntimeError("Lipsesc SUPABASE_URL sau SUPABASE_ANON_KEY în Streamlit Secrets.")
    return create_client(url, key)

def auth_user_id() -> str | None:
    user = st.session_state.get("auth_user")
    return getattr(user, "id", None) if user else None

def current_user_email() -> str:
    user = st.session_state.get("auth_user")
    return getattr(user, "email", "") or ""

def sign_in(email: str, password: str) -> None:
    response = get_supabase().auth.sign_in_with_password({"email": email, "password": password})
    st.session_state["auth_user"] = response.user
    st.session_state["auth_session"] = response.session

def sign_up(email: str, password: str) -> str:
    response = get_supabase().auth.sign_up({"email": email, "password": password})
    if response.session:
        st.session_state["auth_user"] = response.user
        st.session_state["auth_session"] = response.session
        return "Cont creat și autentificat."
    return "Cont creat. Verifică emailul, apoi autentifică-te."

def sign_out() -> None:
    try:
        get_supabase().auth.sign_out()
    finally:
        st.session_state.clear()

def require_login() -> None:
    if auth_user_id():
        return
    st.title("🇪🇺 GrantAI Europe — Etapa 8")
    st.caption("Funding Opportunity Engine")
    login_tab, signup_tab = st.tabs(["Autentificare", "Cont nou"])
    with login_tab:
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Parolă", type="password", key="login_password")
        if st.button("Intră în cont", type="primary"):
            try:
                sign_in(email.strip(), password)
                st.rerun()
            except Exception as exc:
                st.error(f"Autentificarea a eșuat: {exc}")
    with signup_tab:
        email = st.text_input("Email", key="signup_email")
        password = st.text_input("Parolă", type="password", key="signup_password")
        confirm = st.text_input("Confirmă parola", type="password")
        if st.button("Creează cont"):
            if len(password) < 8:
                st.error("Parola trebuie să aibă cel puțin 8 caractere.")
            elif password != confirm:
                st.error("Parolele nu coincid.")
            else:
                try:
                    st.success(sign_up(email.strip(), password))
                except Exception as exc:
                    st.error(f"Crearea contului a eșuat: {exc}")
    st.stop()

def load_organisation() -> dict[str, Any]:
    response = get_supabase().table("organisations").select("data").limit(1).execute()
    if response.data:
        return response.data[0]["data"]
    get_supabase().table("organisations").insert({"user_id": auth_user_id(), "data": DEFAULT_ORG}).execute()
    return json.loads(json.dumps(DEFAULT_ORG))

def save_organisation(data: dict[str, Any]) -> None:
    existing = get_supabase().table("organisations").select("id").limit(1).execute()
    payload = {
        "user_id": auth_user_id(),
        "data": data,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if existing.data:
        get_supabase().table("organisations").update(payload).eq("id", existing.data[0]["id"]).execute()
    else:
        get_supabase().table("organisations").insert(payload).execute()

def list_projects() -> list[dict[str, Any]]:
    response = get_supabase().table("projects").select("id,data,updated_at").order("updated_at", desc=True).execute()
    return [{"id": row["id"], **row["data"]} for row in response.data]

def create_default_project_if_needed() -> None:
    if not list_projects():
        get_supabase().table("projects").insert({
            "user_id": auth_user_id(),
            "name": DEFAULT_PROJECT["name"],
            "data": DEFAULT_PROJECT,
        }).execute()

def save_project(project_id: str | None, project: dict[str, Any]) -> None:
    payload = {
        "user_id": auth_user_id(),
        "name": project["name"],
        "data": project,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if project_id:
        get_supabase().table("projects").update(payload).eq("id", project_id).execute()
    else:
        get_supabase().table("projects").insert(payload).execute()

def save_opportunity(item: dict[str, Any], auto_saved: bool = False) -> None:
    identity = item.get("id") or item.get("reference") or item.get("title")
    existing = (
        get_supabase()
        .table("opportunities")
        .select("id")
        .eq("identity", identity)
        .limit(1)
        .execute()
    )
    item_to_store = dict(item)
    item_to_store["auto_saved"] = auto_saved
    payload = {
        "user_id": auth_user_id(),
        "identity": identity,
        "data": item_to_store,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }
    if existing.data:
        get_supabase().table("opportunities").update(payload).eq("id", existing.data[0]["id"]).execute()
    else:
        get_supabase().table("opportunities").insert(payload).execute()

def list_opportunities() -> list[dict[str, Any]]:
    response = get_supabase().table("opportunities").select("data").order("updated_at", desc=True).execute()
    return [row["data"] for row in response.data]

def save_analysis(project_id: str, opportunity: dict[str, Any], content: str) -> None:
    get_supabase().table("analyses").insert({
        "user_id": auth_user_id(),
        "project_id": project_id,
        "opportunity_identity": opportunity.get("id") or opportunity.get("reference"),
        "content": content,
    }).execute()

def save_document(project_id: str, opportunity: dict[str, Any], title: str, content: str) -> None:
    get_supabase().table("documents").insert({
        "user_id": auth_user_id(),
        "project_id": project_id,
        "opportunity_identity": opportunity.get("id") or opportunity.get("reference"),
        "title": title,
        "content": content,
    }).execute()

def list_history(table: str) -> list[dict[str, Any]]:
    response = get_supabase().table(table).select("*").order("created_at", desc=True).execute()
    return response.data

def record_sync(query: str, count: int) -> None:
    get_supabase().table("sync_runs").insert({
        "user_id": auth_user_id(),
        "query": query,
        "result_count": count,
    }).execute()

def list_sync_runs() -> list[dict[str, Any]]:
    return list_history("sync_runs")

def clean_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return ", ".join(str(x) for x in value)
    if isinstance(value, dict):
        return ", ".join(str(x) for x in value.values())
    return str(value)

def official_search_url(reference: str, title: str) -> str:
    query = reference.strip() or title.strip()
    return f"{PORTAL_BASE}?keywords={quote_plus(query)}"

def normalize_result(item: dict[str, Any]) -> dict[str, Any]:
    source = item.get("_source", item)
    reference = clean_value(source.get("reference"))
    title = clean_value(source.get("title"))
    return {
        "id": clean_value(source.get("identifier") or source.get("callccm2Id") or reference),
        "reference": reference,
        "title": title,
        "status": clean_value(source.get("status")),
        "programme": clean_value(source.get("frameworkProgramme") or source.get("programme") or source.get("caName")),
        "action_type": clean_value(source.get("typesOfAction") or source.get("typeOfAction")),
        "opening_date": clean_value(source.get("startDate")),
        "deadline": clean_value(source.get("deadlineDate") or source.get("deadline")),
        "description": clean_value(source.get("description")),
        "official_url": official_search_url(reference, title),
    }

@st.cache_data(ttl=1800, show_spinner=False)
def search_eu_calls(keyword: str, page_size: int) -> list[dict[str, Any]]:
    query = {
        "bool": {
            "must": [
                {"terms": {"type": ["1", "2", "8"]}},
                {"term": {"programmePeriod": "2021 - 2027"}},
            ]
        }
    }
    params = {
        "apiKey": EC_API_KEY,
        "text": keyword.strip() or "***",
        "pageSize": str(page_size),
        "pageNumber": "1",
    }
    files = {
        "query": ("query.json", json.dumps(query), "application/json"),
        "sort": ("sort.json", json.dumps({"order": "ASC", "field": "deadlineDate"}), "application/json"),
        "languages": ("languages.json", json.dumps(["en"]), "application/json"),
        "displayFields": ("fields.json", json.dumps([
            "type","identifier","reference","callccm2Id","title","status",
            "caName","startDate","deadlineDate","frameworkProgramme",
            "typesOfAction","description"
        ]), "application/json"),
    }
    headers = {
        "Accept": "application/json",
        "Origin": "https://ec.europa.eu",
        "Referer": "https://ec.europa.eu/",
        "X-Requested-With": "XMLHttpRequest",
    }
    response = requests.post(EC_SEARCH_API, params=params, files=files, headers=headers, timeout=60)
    response.raise_for_status()
    payload = response.json()
    candidates = (
        payload.get("results")
        or payload.get("hits", {}).get("hits")
        or payload.get("response", {}).get("docs")
        or []
    )
    return [normalize_result(x) for x in candidates]

def overlap_score(words: list[str], text: str) -> float:
    if not words:
        return 0.0
    corpus = text.lower()
    hits = sum(1 for word in words if word.strip().lower() in corpus)
    return min(100.0, hits / len(words) * 100.0)

def deadline_info(value: str) -> tuple[float, str, int | None]:
    if not value:
        return 45.0, "Necunoscut", None
    try:
        deadline = date_parser.parse(value)
        if deadline.tzinfo is None:
            deadline = deadline.replace(tzinfo=timezone.utc)
        days = (deadline - datetime.now(timezone.utc)).days
        if days < 0:
            return 0.0, "Închis", days
        if days < 14:
            return 15.0, f"{days} zile", days
        if days < 30:
            return 35.0, f"{days} zile", days
        if days < 60:
            return 60.0, f"{days} zile", days
        if days < 120:
            return 82.0, f"{days} zile", days
        return 100.0, f"{days} zile", days
    except Exception:
        return 40.0, "Format necunoscut", None

def score_opportunity(call: dict[str, Any], org: dict[str, Any], project: dict[str, Any]) -> dict[str, Any]:
    corpus = " ".join([
        call.get("title", ""),
        call.get("description", ""),
        call.get("programme", ""),
        call.get("action_type", ""),
    ])
    thematic = overlap_score(project["keywords"], corpus)
    capability = overlap_score(org["capabilities"], corpus)
    timing, deadline_label, days = deadline_info(call.get("deadline", ""))
    evidence = 80.0 if call.get("description") else 45.0
    total = thematic * 0.38 + capability * 0.24 + timing * 0.14 + 70 * 0.16 + evidence * 0.08
    return {
        **call,
        "score": round(max(0, min(100, total)), 1),
        "thematic_fit": round(thematic, 1),
        "capability_fit": round(capability, 1),
        "deadline_label": deadline_label,
        "days_left": days,
    }

def filter_opportunities(items, programme, minimum_score, minimum_days, only_open):
    output = []
    for item in items:
        if programme != "Toate" and programme.lower() not in item.get("programme", "").lower():
            continue
        if item.get("score", 0) < minimum_score:
            continue
        days = item.get("days_left")
        if only_open and days is not None and days < 0:
            continue
        if days is not None and days < minimum_days:
            continue
        output.append(item)
    return output

def ai_generate(task, opportunity, org, project) -> str:
    key = get_secret("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("Lipsește OPENAI_API_KEY în Streamlit Secrets.")
    model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")
    client = OpenAI(api_key=key)
    response = client.responses.create(
        model=model,
        instructions=(
            "Ești consultant senior și evaluator pentru finanțări europene. "
            "Scrie în română. Nu inventa condiții, parteneri, bugete, TRL, "
            "certificări, experiență sau rezultate. Pentru orice lipsă scrie [DE COMPLETAT]."
        ),
        input=f"""
SARCINĂ:
{task}

ORGANIZAȚIE:
{json.dumps(org, ensure_ascii=False, indent=2)}

PROIECT:
{json.dumps(project, ensure_ascii=False, indent=2)}

APEL:
{json.dumps(opportunity, ensure_ascii=False, indent=2)}
""",
    )
    return response.output_text

def markdown_to_docx(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()

require_login()
create_default_project_if_needed()

org = load_organisation()
projects = list_projects()

st.sidebar.write(f"**Cont:** {current_user_email()}")
if st.sidebar.button("Ieșire din cont"):
    sign_out()
    st.rerun()

project_map = {f'{p["name"]} — {p["id"][:8]}': p for p in projects}
selected_project_label = st.sidebar.selectbox("Proiect activ", list(project_map.keys()))
active_project = project_map[selected_project_label]
active_project_id = active_project["id"]

st.title("🇪🇺 GrantAI Europe — Etapa 8")
st.caption("Funding Opportunity Engine")

tabs = st.tabs([
    "Dashboard","Organizație","Proiecte","Opportunity Engine",
    "Selectate","Analiză AI","Generator","Istoric"
])

dashboard_tab, org_tab, projects_tab, engine_tab, saved_tab, analysis_tab, generator_tab, history_tab = tabs

with dashboard_tab:
    opportunities = list_opportunities()
    analyses = list_history("analyses")
    documents = list_history("documents")
    sync_runs = list_sync_runs()

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Proiecte", len(projects))
    c2.metric("Oportunități", len(opportunities))
    c3.metric("Analize", len(analyses))
    c4.metric("Documente", len(documents))

    if sync_runs:
        latest = sync_runs[0]
        st.info(f'Ultima sincronizare: {latest["created_at"]} · {latest["result_count"]} rezultate')
    else:
        st.info("Nu există sincronizări încă.")

with org_tab:
    st.subheader("Organizație")
    c1, c2 = st.columns(2)
    org["legal_name"] = c1.text_input("Denumire legală", org["legal_name"])
    org["country"] = c2.text_input("Țară", org["country"])
    org["organisation_type"] = c1.text_input("Tip organizație", org["organisation_type"])
    org["pic"] = c2.text_input("PIC", org["pic"])
    org["caen"] = c1.text_input("CAEN", org["caen"])
    org["staff"] = c2.number_input("Angajați", min_value=0, value=int(org["staff"]))
    org["turnover_eur"] = c1.number_input("Cifră de afaceri (€)", min_value=0, value=int(org["turnover_eur"]))
    org["capabilities"] = [
        x.strip() for x in st.text_area("Capabilități", ", ".join(org["capabilities"])).split(",") if x.strip()
    ]
    if st.button("Salvează organizația", type="primary"):
        save_organisation(org)
        st.success("Organizație salvată.")

with projects_tab:
    st.subheader("Proiect activ")
    project = {k: v for k, v in active_project.items() if k != "id"}
    project["name"] = st.text_input("Nume", project["name"])
    project["summary"] = st.text_area("Rezumat", project["summary"], height=130)
    project["keywords"] = [
        x.strip() for x in st.text_input("Cuvinte-cheie", ", ".join(project["keywords"])).split(",") if x.strip()
    ]
    project["target_budget_eur"] = st.number_input("Buget țintă (€)", min_value=0, value=int(project["target_budget_eur"]))
    roles = ["beneficiary", "partner", "coordinator", "subcontractor"]
    role = project.get("preferred_role", "beneficiary")
    project["preferred_role"] = st.selectbox("Rol", roles, index=roles.index(role) if role in roles else 0)
    if st.button("Salvează proiectul", type="primary"):
        save_project(active_project_id, project)
        st.success("Proiect salvat.")
        st.rerun()

with engine_tab:
    st.subheader("Funding Opportunity Engine")
    c1, c2 = st.columns([3, 1])
    query = c1.text_input("Căutare automată", " ".join(active_project.get("keywords", [])))
    api_limit = c2.selectbox("Rezultate API", [20, 50, 100], index=1)

    f1, f2, f3, f4 = st.columns(4)
    programme = f1.selectbox("Program", PROGRAMMES)
    minimum_score = f2.slider("Scor minim", 0, 100, 45, 5)
    minimum_days = f3.slider("Minimum zile rămase", 0, 180, 21, 7)
    only_open = f4.checkbox("Doar active", value=True)
    auto_save_threshold = st.slider("Salvează automat peste scorul", 50, 100, 75, 5)

    if st.button("Sincronizează și evaluează", type="primary", use_container_width=True):
        with st.spinner("Preiau, deduplic și evaluez apelurile..."):
            try:
                raw = search_eu_calls(query, api_limit)
                unique = {}
                for item in raw:
                    identity = item.get("id") or item.get("reference") or item.get("title")
                    if identity:
                        unique[identity] = item
                scored = [score_opportunity(item, org, active_project) for item in unique.values()]
                scored.sort(key=lambda x: x["score"], reverse=True)
                st.session_state["engine_results"] = scored
                record_sync(query, len(scored))

                auto_saved = 0
                for item in scored:
                    if item["score"] >= auto_save_threshold:
                        save_opportunity(item, auto_saved=True)
                        auto_saved += 1

                st.success(f"{len(scored)} apeluri evaluate. {auto_saved} salvate automat.")
            except Exception as exc:
                st.error(f"Sincronizarea a eșuat: {exc}")

    results = st.session_state.get("engine_results", [])
    filtered = filter_opportunities(results, programme, minimum_score, minimum_days, only_open)

    if filtered:
        st.write(f"**Oportunități după filtre: {len(filtered)}**")
        st.dataframe(pd.DataFrame([{
            "Scor": x["score"],
            "Referință": x["reference"],
            "Titlu": x["title"],
            "Program": x["programme"],
            "Tip": x["action_type"],
            "Deadline": x["deadline"],
            "Zile": x["days_left"],
        } for x in filtered]), use_container_width=True, hide_index=True)

        idx = st.selectbox(
            "Deschide oportunitatea",
            range(len(filtered)),
            format_func=lambda i: f'{filtered[i]["score"]}% — {filtered[i]["reference"]} — {filtered[i]["title"]}',
        )
        selected = filtered[idx]

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Scor total", f'{selected["score"]}%')
        m2.metric("Tematic", f'{selected["thematic_fit"]}%')
        m3.metric("Capabilități", f'{selected["capability_fit"]}%')
        m4.metric("Timp", selected["deadline_label"])

        st.write(f'**Program:** {selected["programme"] or "N/A"}')
        st.write(f'**Tip acțiune:** {selected["action_type"] or "N/A"}')
        st.write(f'**Deadline:** {selected["deadline"] or "N/A"}')
        st.link_button("Deschide portalul oficial", selected["official_url"])

        if selected["description"]:
            with st.expander("Descriere publică"):
                st.write(selected["description"])

        if st.button("Salvează manual"):
            save_opportunity(selected, auto_saved=False)
            st.success("Oportunitate salvată.")
    elif results:
        st.warning("Niciun apel nu respectă filtrele.")

with saved_tab:
    opportunities = list_opportunities()
    st.subheader("Oportunități salvate")
    if not opportunities:
        st.info("Nu există oportunități.")
    else:
        st.dataframe(pd.DataFrame([{
            "Scor": x.get("score"),
            "Referință": x.get("reference"),
            "Titlu": x.get("title"),
            "Deadline": x.get("deadline"),
            "Automat": "Da" if x.get("auto_saved") else "Nu",
        } for x in opportunities]), use_container_width=True, hide_index=True)

with analysis_tab:
    opportunities = list_opportunities()
    st.subheader("Analiză AI")
    if not opportunities:
        st.warning("Salvează o oportunitate.")
    else:
        idx = st.selectbox(
            "Oportunitate",
            range(len(opportunities)),
            format_func=lambda i: f'{opportunities[i].get("reference")} — {opportunities[i].get("title")}',
            key="analysis_opportunity",
        )
        opportunity = opportunities[idx]
        if st.button("Analizează cu AI", type="primary"):
            task = """
Realizează verdict GO / CONDITIONAL GO / NO-GO, potrivire, eligibilitate,
parteneri lipsă, documente oficiale necesare, riscuri, plan pe 14 zile și scor 0-100.
"""
            with st.spinner("AI analizează..."):
                try:
                    result = ai_generate(task, opportunity, org, active_project)
                    save_analysis(active_project_id, opportunity, result)
                    st.session_state["analysis_result"] = result
                except Exception as exc:
                    st.error(str(exc))

        result = st.session_state.get("analysis_result")
        if result:
            st.markdown(result)
            st.download_button(
                "Descarcă Word",
                markdown_to_docx("Analiză GrantAI", result),
                file_name="analiza_grantai.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

with generator_tab:
    opportunities = list_opportunities()
    st.subheader("Generator")
    if not opportunities:
        st.warning("Salvează o oportunitate.")
    else:
        idx = st.selectbox(
            "Oportunitate",
            range(len(opportunities)),
            format_func=lambda i: f'{opportunities[i].get("reference")} — {opportunities[i].get("title")}',
            key="generator_opportunity",
        )
        opportunity = opportunities[idx]
        doc_type = st.selectbox(
            "Document",
            [
                "Concept note","Excellence","Impact","Implementation","Work packages",
                "Deliverables și milestones","Plan de impact și KPI","Registru de riscuri","Rezumat executiv"
            ],
        )
        if st.button("Generează documentul", type="primary"):
            task = f"""
Generează {doc_type}. Folosește numai datele disponibile.
Pentru lipsuri scrie [DE COMPLETAT].
Include tabel cerință / răspuns / dovadă și checklist final.
"""
            with st.spinner("AI redactează..."):
                try:
                    result = ai_generate(task, opportunity, org, active_project)
                    save_document(active_project_id, opportunity, doc_type, result)
                    st.session_state["generated_result"] = result
                except Exception as exc:
                    st.error(str(exc))

        result = st.session_state.get("generated_result")
        if result:
            st.markdown(result)
            st.download_button(
                "Descarcă Word",
                markdown_to_docx(doc_type, result),
                file_name=f'{doc_type.lower().replace(" ", "_")}.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

with history_tab:
    st.subheader("Istoric")
    sync_runs = list_sync_runs()
    analyses = list_history("analyses")
    documents = list_history("documents")

    st.write("### Sincronizări")
    if sync_runs:
        st.dataframe(pd.DataFrame([{
            "Data": x["created_at"],
            "Query": x["query"],
            "Rezultate": x["result_count"],
        } for x in sync_runs]), use_container_width=True, hide_index=True)
    else:
        st.info("Nu există sincronizări.")

    st.write("### Analize AI")
    for row in analyses:
        with st.expander(f'Analiză — {row["created_at"]}'):
            st.markdown(row["content"])

    st.write("### Documente")
    for row in documents:
        with st.expander(f'{row["title"]} — {row["created_at"]}'):
            st.markdown(row["content"])

st.divider()
st.caption(
    "Funding Opportunity Engine folosește date publice Funding & Tenders. "
    "Depunerea oficială și EU Login rămân în portalul Comisiei Europene."
)
