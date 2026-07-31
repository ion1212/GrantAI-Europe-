import io
import json
import os
from datetime import datetime, timezone
from typing import Any

import streamlit as st
from docx import Document
from openai import OpenAI
from supabase import create_client

st.set_page_config(page_title="AI Proposal Generator", page_icon="📝", layout="wide")

DOCUMENT_TYPES = [
    "Executive Summary", "Abstract", "Concept Note", "Excellence", "Impact",
    "Implementation", "Work Packages", "Deliverables and Milestones",
    "Risk Management", "Dissemination Plan", "Exploitation Plan", "Ethics",
    "Gender Dimension", "Budget Narrative",
]

CHECKLIST_ITEMS = [
    "Excellence", "Impact", "Implementation", "Ethics",
    "Budget", "Consortium", "Risks", "Dissemination",
]

def secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default))
    except Exception:
        return os.getenv(name, default)

@st.cache_resource
def supabase():
    url = secret("SUPABASE_URL")
    key = secret("SUPABASE_ANON_KEY")
    if not url or not key:
        raise RuntimeError("Lipsesc SUPABASE_URL sau SUPABASE_ANON_KEY.")
    return create_client(url, key)

def user_id() -> str | None:
    user = st.session_state.get("auth_user")
    return getattr(user, "id", None) if user else None

def require_login() -> None:
    if not user_id():
        st.error("Autentifică-te mai întâi în pagina principală GrantAI Europe.")
        st.stop()

def load_org() -> dict[str, Any]:
    result = supabase().table("organisations").select("data").limit(1).execute()
    return result.data[0]["data"] if result.data else {}

def load_projects() -> list[dict[str, Any]]:
    result = (
        supabase().table("projects")
        .select("id,data,updated_at")
        .order("updated_at", desc=True)
        .execute()
    )
    return [{"id": row["id"], **row["data"]} for row in result.data]

def load_opportunities() -> list[dict[str, Any]]:
    result = (
        supabase().table("opportunities")
        .select("data")
        .order("updated_at", desc=True)
        .execute()
    )
    return [row["data"] for row in result.data]

def opportunity_identity(opportunity: dict[str, Any]) -> str:
    return str(
        opportunity.get("id")
        or opportunity.get("reference")
        or opportunity.get("title")
        or ""
    )

def latest_analysis(project_id: str, identity: str) -> str:
    result = (
        supabase().table("analyses")
        .select("content")
        .eq("project_id", project_id)
        .eq("opportunity_identity", identity)
        .order("created_at", desc=True)
        .limit(1)
        .execute()
    )
    return result.data[0]["content"] if result.data else ""

def next_version(project_id: str, identity: str, document_type: str) -> int:
    result = (
        supabase().table("proposal_versions")
        .select("version")
        .eq("project_id", project_id)
        .eq("opportunity_identity", identity)
        .eq("document_type", document_type)
        .order("version", desc=True)
        .limit(1)
        .execute()
    )
    return int(result.data[0]["version"]) + 1 if result.data else 1

def save_version(
    project_id: str,
    identity: str,
    document_type: str,
    language: str,
    content: str,
    quality_score: int | None = None,
    checklist: dict[str, bool] | None = None,
) -> int:
    version = next_version(project_id, identity, document_type)
    supabase().table("proposal_versions").insert({
        "user_id": user_id(),
        "project_id": project_id,
        "opportunity_identity": identity,
        "document_type": document_type,
        "language": language,
        "version": version,
        "content": content,
        "quality_score": quality_score,
        "checklist": checklist or {},
        "created_at": datetime.now(timezone.utc).isoformat(),
    }).execute()
    return version

def load_versions(project_id: str) -> list[dict[str, Any]]:
    return (
        supabase().table("proposal_versions")
        .select("*")
        .eq("project_id", project_id)
        .order("created_at", desc=True)
        .execute()
        .data
    )

def openai_client() -> OpenAI:
    key = secret("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("Lipsește OPENAI_API_KEY.")
    return OpenAI(api_key=key)

def generate_document(
    organisation: dict[str, Any],
    project: dict[str, Any],
    opportunity: dict[str, Any],
    analysis: str,
    document_type: str,
    language: str,
    instructions: str,
) -> str:
    model = secret("OPENAI_MODEL", "gpt-4.1-mini")
    lang = "Write in English." if language == "English" else "Scrie în limba română."
    prompt = f"""
You are a senior EU funding proposal writer and evaluator.
{lang}

Generate a proposal-ready draft for: {document_type}.

Mandatory rules:
- Do not invent eligibility conditions, budget ceilings, TRL, consortium members,
  certifications, past performance, results or official requirements.
- Mark missing information as [TO COMPLETE] in English or [DE COMPLETAT] in Romanian.
- Separate confirmed facts from assumptions.
- Use clear headings and measurable language.
- End with a validation checklist.
- This is a draft, not confirmation of compliance.

ORGANISATION:
{json.dumps(organisation, ensure_ascii=False, indent=2)}

PROJECT:
{json.dumps(project, ensure_ascii=False, indent=2)}

OPPORTUNITY:
{json.dumps(opportunity, ensure_ascii=False, indent=2)}

LATEST ANALYSIS:
{analysis or "[No prior analysis available]"}

ADDITIONAL INSTRUCTIONS:
{instructions or "[None]"}
"""
    response = openai_client().responses.create(model=model, input=prompt)
    return response.output_text

def evaluate_document(
    content: str,
    document_type: str,
    opportunity: dict[str, Any],
    language: str,
) -> dict[str, Any]:
    model = secret("OPENAI_MODEL", "gpt-4.1-mini")
    prompt = f"""
Evaluate this EU proposal draft.
Return only valid JSON, without markdown fences.

Required schema:
{{
  "quality_score": 0,
  "missing_sections": [],
  "contradictions": [],
  "unsupported_claims": [],
  "recommendations": [],
  "checklist": {{
    "Excellence": false,
    "Impact": false,
    "Implementation": false,
    "Ethics": false,
    "Budget": false,
    "Consortium": false,
    "Risks": false,
    "Dissemination": false
  }}
}}

DOCUMENT TYPE: {document_type}
LANGUAGE: {language}

OPPORTUNITY:
{json.dumps(opportunity, ensure_ascii=False, indent=2)}

DRAFT:
{content}
"""
    raw = openai_client().responses.create(model=model, input=prompt).output_text.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {
            "quality_score": 0,
            "missing_sections": ["Răspunsul evaluatorului nu a putut fi interpretat."],
            "contradictions": [],
            "unsupported_claims": [],
            "recommendations": [raw],
            "checklist": {item: False for item in CHECKLIST_ITEMS},
        }

def to_docx(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for raw in content.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

require_login()

organisation = load_org()
projects = load_projects()
opportunities = load_opportunities()

st.title("📝 Etapa 10 — AI Proposal Generator")
st.caption("Generator pe secțiuni, editor, evaluare și versiuni în Supabase")

if not projects:
    st.error("Nu există proiecte.")
    st.stop()

project_labels = {
    f'{project.get("name", "Proiect")} — {project["id"][:8]}': project
    for project in projects
}
project_label = st.selectbox("Proiect", list(project_labels))
project = project_labels[project_label]
project_id = project["id"]

generator_tab, editor_tab, evaluator_tab, versions_tab = st.tabs([
    "Generator", "Editor", "Evaluator", "Versiuni"
])

with generator_tab:
    if not opportunities:
        st.warning(
            "Nu există oportunități salvate. Revino în Opportunity Engine și "
            "apasă «Salvează manual» pentru un apel."
        )
    else:
        opportunity_index = st.selectbox(
            "Oportunitate",
            range(len(opportunities)),
            format_func=lambda i: (
                f'{opportunities[i].get("reference") or "fără cod"} — '
                f'{opportunities[i].get("title") or "fără titlu"}'
            ),
        )
        opportunity = opportunities[opportunity_index]
        identity = opportunity_identity(opportunity)

        col1, col2 = st.columns(2)
        document_type = col1.selectbox("Tip document", DOCUMENT_TYPES)
        language = col2.selectbox("Limba", ["English", "Română"])
        instructions = st.text_area(
            "Instrucțiuni suplimentare",
            placeholder="Ex.: accent pe impact rural, eficiență energetică și replicare.",
            height=100,
        )

        analysis = latest_analysis(project_id, identity)
        if analysis:
            with st.expander("Analiza AI inclusă în context"):
                st.markdown(analysis)

        if st.button("Generează draftul", type="primary", use_container_width=True):
            with st.spinner("AI generează documentul..."):
                try:
                    content = generate_document(
                        organisation,
                        project,
                        opportunity,
                        analysis,
                        document_type,
                        language,
                        instructions,
                    )
                    st.session_state["stage10_content"] = content
                    st.session_state["stage10_document_type"] = document_type
                    st.session_state["stage10_language"] = language
                    st.session_state["stage10_opportunity"] = opportunity
                    st.session_state["stage10_identity"] = identity
                    st.session_state["stage10_evaluation"] = None
                    st.success("Draft generat. Deschide fila Editor.")
                except Exception as exc:
                    st.error(str(exc))

with editor_tab:
    content = st.session_state.get("stage10_content", "")
    if not content:
        st.info("Generează mai întâi un document.")
    else:
        edited = st.text_area(
            "Editor",
            value=content,
            height=650,
            key="stage10_editor",
        )
        st.session_state["stage10_content"] = edited

        col1, col2, col3 = st.columns(3)
        if col1.button("Salvează versiune", type="primary"):
            version = save_version(
                project_id,
                st.session_state["stage10_identity"],
                st.session_state["stage10_document_type"],
                st.session_state["stage10_language"],
                edited,
            )
            st.success(f"Versiunea {version} a fost salvată.")

        col2.download_button(
            "Descarcă Word",
            to_docx(st.session_state["stage10_document_type"], edited),
            file_name=st.session_state["stage10_document_type"].lower().replace(" ", "_") + ".docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        col3.download_button(
            "Descarcă Markdown",
            edited.encode("utf-8"),
            file_name=st.session_state["stage10_document_type"].lower().replace(" ", "_") + ".md",
            mime="text/markdown",
        )

with evaluator_tab:
    content = st.session_state.get("stage10_content", "")
    opportunity = st.session_state.get("stage10_opportunity")

    if not content or not opportunity:
        st.info("Generează și editează mai întâi un document.")
    else:
        if st.button("Evaluează draftul cu AI", type="primary"):
            with st.spinner("AI verifică draftul..."):
                try:
                    st.session_state["stage10_evaluation"] = evaluate_document(
                        content,
                        st.session_state["stage10_document_type"],
                        opportunity,
                        st.session_state["stage10_language"],
                    )
                except Exception as exc:
                    st.error(str(exc))

        evaluation = st.session_state.get("stage10_evaluation")
        if evaluation:
            score = max(0, min(100, int(evaluation.get("quality_score", 0))))
            st.metric("Scor estimativ", f"{score}/100")
            st.progress(score)

            checklist = evaluation.get("checklist", {})
            columns = st.columns(4)
            for index, item in enumerate(CHECKLIST_ITEMS):
                columns[index % 4].write(
                    f'{"✅" if checklist.get(item, False) else "⬜"} {item}'
                )

            for label, key in [
                ("Secțiuni lipsă", "missing_sections"),
                ("Contradicții", "contradictions"),
                ("Afirmații fără suport", "unsupported_claims"),
                ("Recomandări", "recommendations"),
            ]:
                with st.expander(label):
                    values = evaluation.get(key, [])
                    if values:
                        for value in values:
                            st.write(f"- {value}")
                    else:
                        st.write("Nimic identificat.")

            if st.button("Salvează versiunea evaluată"):
                version = save_version(
                    project_id,
                    st.session_state["stage10_identity"],
                    st.session_state["stage10_document_type"],
                    st.session_state["stage10_language"],
                    content,
                    quality_score=score,
                    checklist=checklist,
                )
                st.success(f"Versiunea evaluată {version} a fost salvată.")

with versions_tab:
    versions = load_versions(project_id)
    if not versions:
        st.info("Nu există versiuni salvate.")
    else:
        for version in versions:
            title = (
                f'{version["document_type"]} · v{version["version"]} · '
                f'{version["created_at"]}'
            )
            with st.expander(title):
                if version.get("quality_score") is not None:
                    st.write(f'**Scor:** {version["quality_score"]}/100')
                st.text_area(
                    "Conținut",
                    value=version["content"],
                    height=350,
                    disabled=True,
                    key=f'version_{version["id"]}',
                )
                st.download_button(
                    "Descarcă Word",
                    to_docx(version["document_type"], version["content"]),
                    file_name=(
                        version["document_type"].lower().replace(" ", "_")
                        + f'_v{version["version"]}.docx'
                    ),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f'download_{version["id"]}',
                )

st.divider()
st.caption(
    "Documentele sunt drafturi. Cerințele, eligibilitatea și bugetul trebuie "
    "verificate în documentația oficială a apelului."
)
