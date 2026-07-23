import io
import json
import os
from datetime import datetime, timezone
from typing import Any

import pandas as pd
import requests
import streamlit as st
from dateutil import parser as date_parser
from docx import Document
from openai import OpenAI

st.set_page_config(page_title="GrantAI Europe", page_icon="🇪🇺", layout="wide")

EC_SEARCH_API = "https://api.tech.ec.europa.eu/search-api/prod/rest/search"
EC_API_KEY = "SEDIA"

DEFAULT_PROFILE = {
    "organisation": {
        "legal_name": "II Ciobotaru Viorel Razvan Ionut",
        "country": "Romania",
        "organisation_type": "SME / întreprindere individuală",
        "pic": "",
        "caen": "0111",
        "turnover_eur": 0,
        "staff": 0,
        "capabilities": [
            "agriculture",
            "smart greenhouse",
            "renewable energy",
            "battery storage",
            "AI automation",
        ],
        "past_projects": [],
    },
    "project": {
        "name": "GreenRise",
        "summary": (
            "Seră inteligentă cu energie regenerabilă, baterii "
            "și automatizare AI pentru producție agricolă."
        ),
        "keywords": [
            "agriculture",
            "greenhouse",
            "energy",
            "battery",
            "AI",
            "rural",
            "agrifood",
        ],
        "target_budget_eur": 0,
        "preferred_role": "beneficiary",
    },
}

def get_secret(name: str, default: str = "") -> str:
    try:
        return str(st.secrets.get(name, default))
    except Exception:
        return os.getenv(name, default)

def ensure_state() -> None:
    if "profile" not in st.session_state:
        st.session_state["profile"] = DEFAULT_PROFILE.copy()
    if "saved_opportunities" not in st.session_state:
        st.session_state["saved_opportunities"] = []
    if "generated_documents" not in st.session_state:
        st.session_state["generated_documents"] = []

def clean_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        return ", ".join(str(x) for x in value)
    if isinstance(value, dict):
        return ", ".join(str(x) for x in value.values())
    return str(value)

def normalize_result(item: dict[str, Any]) -> dict[str, Any]:
    source = item.get("_source", item)
    return {
        "id": clean_value(source.get("identifier") or source.get("callccm2Id") or source.get("reference")),
        "reference": clean_value(source.get("reference")),
        "title": clean_value(source.get("title")),
        "status": clean_value(source.get("status")),
        "programme": clean_value(
            source.get("frameworkProgramme")
            or source.get("programme")
            or source.get("caName")
        ),
        "action_type": clean_value(
            source.get("typesOfAction")
            or source.get("typeOfAction")
        ),
        "deadline": clean_value(
            source.get("deadlineDate")
            or source.get("deadline")
        ),
        "description": clean_value(source.get("description")),
    }

@st.cache_data(ttl=3600, show_spinner=False)
def search_eu_calls(keyword: str, page_size: int) -> list[dict[str, Any]]:
    query = {
        "bool": {
            "must": [
                {"terms": {"type": ["1", "2", "8"]}},
                {"term": {"programmePeriod": "2021 - 2027"}},
            ]
        }
    }
    params = {
        "apiKey": EC_API_KEY,
        "text": keyword.strip() or "***",
        "pageSize": str(page_size),
        "pageNumber": "1",
    }
    multipart_files = {
        "query": ("query.json", json.dumps(query), "application/json"),
        "sort": ("sort.json", json.dumps({"order": "ASC", "field": "deadlineDate"}), "application/json"),
        "languages": ("languages.json", json.dumps(["en"]), "application/json"),
        "displayFields": ("fields.json", json.dumps([
            "type", "identifier", "reference", "callccm2Id", "title", "status",
            "caName", "startDate", "deadlineDate", "frameworkProgramme",
            "typesOfAction", "description"
        ]), "application/json"),
    }
    headers = {
        "Accept": "application/json",
        "Origin": "https://ec.europa.eu",
        "Referer": "https://ec.europa.eu/",
        "X-Requested-With": "XMLHttpRequest",
    }
    response = requests.post(
        EC_SEARCH_API,
        params=params,
        files=multipart_files,
        headers=headers,
        timeout=60,
    )
    response.raise_for_status()
    payload = response.json()
    candidates = (
        payload.get("results")
        or payload.get("hits", {}).get("hits")
        or payload.get("response", {}).get("docs")
        or []
    )
    return [normalize_result(x) for x in candidates]

def overlap_score(words: list[str], text: str) -> float:
    if not words:
        return 0.0
    corpus = text.lower()
    hits = sum(1 for word in words if word.strip().lower() in corpus)
    return min(100.0, hits / len(words) * 100.0)

def deadline_score(value: str) -> tuple[float, str]:
    if not value:
        return 45.0, "Necunoscut"
    try:
        deadline = date_parser.parse(value)
        if deadline.tzinfo is None:
            deadline = deadline.replace(tzinfo=timezone.utc)
        days = (deadline - datetime.now(timezone.utc)).days
        if days < 0:
            return 0.0, "Închis"
        if days < 14:
            return 15.0, f"{days} zile"
        if days < 30:
            return 35.0, f"{days} zile"
        if days < 60:
            return 60.0, f"{days} zile"
        if days < 120:
            return 82.0, f"{days} zile"
        return 100.0, f"{days} zile"
    except Exception:
        return 40.0, "Format necunoscut"

def score_opportunity(call: dict[str, Any], profile: dict[str, Any]) -> dict[str, Any]:
    corpus = " ".join([
        call.get("title", ""),
        call.get("description", ""),
        call.get("programme", ""),
        call.get("action_type", ""),
    ])
    thematic = overlap_score(profile["project"]["keywords"], corpus)
    capability = overlap_score(profile["organisation"]["capabilities"], corpus)
    timing, deadline_label = deadline_score(call.get("deadline", ""))
    evidence = 80.0 if call.get("description") else 45.0
    total = thematic * .36 + capability * .24 + timing * .16 + 70 * .16 + evidence * .08
    return {
        **call,
        "score": round(max(0, min(100, total)), 1),
        "thematic_fit": round(thematic, 1),
        "capability_fit": round(capability, 1),
        "deadline_label": deadline_label,
    }

def ai_generate(task: str, opportunity: dict[str, Any], profile: dict[str, Any]) -> str:
    key = get_secret("OPENAI_API_KEY")
    if not key:
        raise RuntimeError("Lipsește OPENAI_API_KEY în Streamlit Secrets.")
    model = get_secret("OPENAI_MODEL", "gpt-4.1-mini")
    client = OpenAI(api_key=key)
    response = client.responses.create(
        model=model,
        instructions=(
            "Ești consultant senior pentru finanțări europene. Scrie în română. "
            "Nu inventa eligibilitate, parteneri, bugete, rezultate, experiență sau certificări. "
            "Pentru informații lipsă scrie [DE COMPLETAT]. "
            "Separă clar faptele confirmate, presupunerile, riscurile și dovezile necesare."
        ),
        input=f"""
SARCINĂ:
{task}

PROFIL:
{json.dumps(profile, ensure_ascii=False, indent=2)}

APEL:
{json.dumps(opportunity, ensure_ascii=False, indent=2)}
""",
    )
    return response.output_text

def markdown_to_docx(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)
    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[:3].rstrip(".").isdigit() and ". " in line:
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        else:
            doc.add_paragraph(line)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

ensure_state()
profile = st.session_state["profile"]

st.title("🇪🇺 GrantAI Europe — Etapa 4")
st.caption("Profil editabil, oportunități reale, analiză AI și export Word")

tabs = st.tabs([
    "Dashboard",
    "Profil",
    "Apeluri",
    "Selectate",
    "Analiză AI",
    "Generator",
    "Documente",
])
tab_dashboard, tab_profile, tab_calls, tab_saved, tab_analysis, tab_generator, tab_docs = tabs

with tab_dashboard:
    st.subheader("Dashboard")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Oportunități selectate", len(st.session_state["saved_opportunities"]))
    c2.metric("Documente generate", len(st.session_state["generated_documents"]))
    c3.metric("Proiect principal", profile["project"]["name"])
    c4.metric("AI", "Activ" if get_secret("OPENAI_API_KEY") else "Neconfigurat")
    st.info(
        "Flux recomandat: completează profilul → caută apeluri → selectează → "
        "rulează analiza AI → generează documentul → descarcă Word."
    )

with tab_profile:
    st.subheader("Profil organizație")
    org = profile["organisation"]
    proj = profile["project"]

    col1, col2 = st.columns(2)
    org["legal_name"] = col1.text_input("Denumire legală", org["legal_name"])
    org["country"] = col2.text_input("Țară", org["country"])
    org["organisation_type"] = col1.text_input("Tip organizație", org["organisation_type"])
    org["pic"] = col2.text_input("PIC", org["pic"])
    org["caen"] = col1.text_input("CAEN", org["caen"])
    org["turnover_eur"] = col2.number_input("Cifră de afaceri estimată (€)", min_value=0, value=int(org["turnover_eur"]))
    org["staff"] = col1.number_input("Număr angajați", min_value=0, value=int(org["staff"]))
    org["capabilities"] = [
        x.strip() for x in st.text_area(
            "Capabilități, separate prin virgulă",
            ", ".join(org["capabilities"]),
        ).split(",") if x.strip()
    ]
    org["past_projects"] = [
        x.strip() for x in st.text_area(
            "Proiecte anterioare, câte unul pe linie",
            "\n".join(org["past_projects"]),
        ).splitlines() if x.strip()
    ]

    st.subheader("Proiect principal")
    proj["name"] = st.text_input("Nume proiect", proj["name"])
    proj["summary"] = st.text_area("Rezumat proiect", proj["summary"], height=150)
    proj["keywords"] = [
        x.strip() for x in st.text_input(
            "Cuvinte-cheie, separate prin virgulă",
            ", ".join(proj["keywords"]),
        ).split(",") if x.strip()
    ]
    proj["target_budget_eur"] = st.number_input(
        "Buget țintă (€)",
        min_value=0,
        value=int(proj["target_budget_eur"]),
    )
    proj["preferred_role"] = st.selectbox(
        "Rol preferat",
        ["beneficiary", "partner", "coordinator", "subcontractor"],
        index=["beneficiary", "partner", "coordinator", "subcontractor"].index(proj["preferred_role"]),
    )

    if st.button("Salvează profilul", type="primary"):
        st.session_state["profile"] = profile
        st.success("Profil salvat în sesiunea curentă.")

    st.download_button(
        "Descarcă profilul JSON",
        data=json.dumps(profile, ensure_ascii=False, indent=2),
        file_name="grantai_profile.json",
        mime="application/json",
    )

with tab_calls:
    st.subheader("Căutare în Funding & Tenders")
    c1, c2 = st.columns([3, 1])
    keyword = c1.text_input("Cuvinte-cheie", "agriculture energy battery AI")
    limit = c2.selectbox("Rezultate", [10, 20, 30, 50], index=1)

    if st.button("Caută apeluri", type="primary", use_container_width=True):
        with st.spinner("Interoghez portalul european..."):
            try:
                results = [score_opportunity(x, profile) for x in search_eu_calls(keyword, limit)]
                results.sort(key=lambda x: x["score"], reverse=True)
                st.session_state["results"] = results
                st.success(f"Au fost analizate {len(results)} oportunități.")
            except Exception as exc:
                st.error(f"Căutarea nu a reușit: {exc}")

    results = st.session_state.get("results", [])
    if results:
        st.dataframe(pd.DataFrame([{
            "Scor": x["score"],
            "Referință": x["reference"],
            "Titlu": x["title"],
            "Program": x["programme"],
            "Deadline": x["deadline"],
        } for x in results]), use_container_width=True, hide_index=True)

        idx = st.selectbox(
            "Deschide oportunitatea",
            range(len(results)),
            format_func=lambda i: f'{results[i]["score"]}% — {results[i]["reference"]} — {results[i]["title"]}'
        )
        selected = results[idx]
        st.subheader(selected["title"] or "Oportunitate")
        st.write(f'**Referință:** {selected["reference"] or "N/A"}')
        st.write(f'**Program:** {selected["programme"] or "N/A"}')
        st.write(f'**Deadline:** {selected["deadline"] or "N/A"}')
        c1, c2, c3 = st.columns(3)
        c1.metric("Scor total", f'{selected["score"]}%')
        c2.metric("Potrivire tematică", f'{selected["thematic_fit"]}%')
        c3.metric("Timp", selected["deadline_label"])

        if selected["description"]:
            with st.expander("Descriere"):
                st.write(selected["description"])

        if st.button("Adaugă la selectate"):
            saved = st.session_state["saved_opportunities"]
            identity = selected.get("id") or selected.get("reference")
            existing = {x.get("id") or x.get("reference") for x in saved}
            if identity not in existing:
                saved.append(selected)
                st.success("Oportunitate adăugată.")
            else:
                st.info("Este deja selectată.")

with tab_saved:
    st.subheader("Oportunități selectate")
    saved = st.session_state["saved_opportunities"]
    if not saved:
        st.info("Nu ai selectat încă nicio oportunitate.")
    else:
        st.dataframe(pd.DataFrame([{
            "Scor": x["score"],
            "Referință": x["reference"],
            "Titlu": x["title"],
            "Deadline": x["deadline"],
        } for x in saved]), use_container_width=True, hide_index=True)
        st.download_button(
            "Descarcă lista JSON",
            data=json.dumps(saved, ensure_ascii=False, indent=2),
            file_name="oportunitati_selectate.json",
            mime="application/json",
        )

with tab_analysis:
    st.subheader("Analiză AI")
    saved = st.session_state["saved_opportunities"]
    if not saved:
        st.warning("Selectează mai întâi o oportunitate.")
    else:
        idx = st.selectbox(
            "Alege oportunitatea",
            range(len(saved)),
            format_func=lambda i: f'{saved[i]["reference"]} — {saved[i]["title"]}',
            key="analysis_call",
        )
        selected = saved[idx]
        if st.button("Rulează analiza AI", type="primary"):
            task = """
Realizează:
1. verdict GO / CONDITIONAL GO / NO-GO;
2. potrivirea tematică;
3. întrebările de eligibilitate;
4. punctele forte;
5. lacunele și riscurile;
6. partenerii necesari;
7. documentele și dovezile lipsă;
8. zece acțiuni următoare;
9. scor de pregătire 0-100, explicat.
"""
            with st.spinner("AI analizează..."):
                try:
                    result = ai_generate(task, selected, profile)
                    st.session_state["analysis_result"] = result
                except Exception as exc:
                    st.error(str(exc))

        result = st.session_state.get("analysis_result")
        if result:
            st.markdown(result)
            st.download_button(
                "Descarcă analiza Word",
                markdown_to_docx("Analiză AI", result),
                file_name="analiza_ai.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

with tab_generator:
    st.subheader("Generator de documente")
    saved = st.session_state["saved_opportunities"]
    if not saved:
        st.warning("Selectează mai întâi o oportunitate.")
    else:
        idx = st.selectbox(
            "Alege oportunitatea",
            range(len(saved)),
            format_func=lambda i: f'{saved[i]["reference"]} — {saved[i]["title"]}',
            key="generator_call",
        )
        selected = saved[idx]
        document_type = st.selectbox(
            "Tip document",
            [
                "Concept note",
                "Excellence",
                "Impact",
                "Implementation",
                "Work packages și deliverables",
                "Plan de impact și KPI",
                "Registru de riscuri",
                "Rezumat executiv",
            ],
        )

        if st.button("Generează cu AI", type="primary"):
            task = f"""
Generează: {document_type}.
Textul trebuie să fie un prim draft profesional și structurat.
Mapează fiecare secțiune la informațiile apelului.
Folosește tabele Markdown unde sunt utile.
Pentru informații nesusținute folosește [DE COMPLETAT].
Încheie cu un checklist factual înainte de utilizarea oficială.
"""
            with st.spinner("AI redactează..."):
                try:
                    result = ai_generate(task, selected, profile)
                    st.session_state["generated_text"] = result
                    st.session_state["generated_documents"].append({
                        "title": document_type,
                        "call": selected["reference"],
                        "content": result,
                    })
                except Exception as exc:
                    st.error(str(exc))

        result = st.session_state.get("generated_text")
        if result:
            st.markdown(result)
            st.download_button(
                "Descarcă Word",
                markdown_to_docx(document_type, result),
                file_name=f'{document_type.lower().replace(" ", "_")}.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.download_button(
                "Descarcă Markdown",
                result,
                file_name=f'{document_type.lower().replace(" ", "_")}.md',
                mime="text/markdown",
            )

with tab_docs:
    st.subheader("Documente generate")
    docs = st.session_state["generated_documents"]
    if not docs:
        st.info("Nu există încă documente generate.")
    else:
        for i, doc in enumerate(reversed(docs), start=1):
            with st.expander(f'{i}. {doc["title"]} — {doc["call"]}'):
                st.markdown(doc["content"])

st.divider()
st.caption(
    "Profilul și documentele sunt păstrate doar în sesiunea curentă. "
    "În etapa următoare vom adăuga bază de date și salvare permanentă."
)
